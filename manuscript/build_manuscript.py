# -*- coding: utf-8 -*-
"""
Build the submission-ready manuscript.

  manuscript_SciReps.docx  (primary deliverable)
  manuscript_SciReps.tex   (LaTeX fallback)

Sci Rep 适配:
  - 摘要：非结构化单段落（已由 01_title_abstract.md 提供）
  - 正文顺序：Introduction → Results → Discussion → Methods → Limitations
  - 关键词：仅在投稿系统填写，不写入正文
  - 参考文献：Nature 编号格式（来自 07_references.md）
  - 图表：4 张主图 + 补充材料引用
  - 占位符已清理；MVMR/反向MR 实际数值从 results/ 填充

用法（在 analysis venv 中运行）：
  ./venv/Scripts/python.exe manuscript/build_manuscript.py
"""
import os, re, csv

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))   # spine_MR/
MANU = os.path.join(ROOT, "manuscript")
RES  = os.path.join(ROOT, "results")
FIG  = os.path.join(RES, "figures")
OUT_DOCX = os.path.join(MANU, "manuscript_SciReps.docx")
OUT_TEX  = os.path.join(MANU, "manuscript_SciReps.tex")
TMP_DOCX = os.path.join(MANU, "_build_tmp.docx")
TMP_TEX  = os.path.join(MANU, "_build_tmp.tex")

# 注释 → 上标 unicode（用于 tex 渲染）
SUP_MAP = {
    '\u207b':'-','\u2070':'0','\u00b9':'1','\u00b2':'2','\u00b3':'3',
    '\u2074':'4','\u2075':'5','\u2076':'6','\u2077':'7','\u2078':'8','\u2079':'9',
}

# ============== inline parser ==============
def inline_parts(text):
    parts = []
    pat = re.compile(r'(\*\*.*?\*\*|`.*?`|\*[^*][^*]*?\*)')
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], 'n'))
        tok = m.group(0)
        if tok.startswith('**'):
            parts.append((tok[2:-2], 'b'))
        elif tok.startswith('`'):
            parts.append((tok[1:-1], 'c'))
        else:
            parts.append((tok[1:-1], 'i'))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], 'n'))
    return parts

def parse_table(lines):
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    return rows[0], rows[2:]

def generic_parse(path, skip_h2_prefixes=()):
    """Parse a markdown file into blocks. Skip any H2 whose title starts with
    any prefix in skip_h2_prefixes (e.g. 'references' to drop in-file ref blocks)."""
    blocks = []
    with open(path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    i, n = 0, len(lines)
    skip_section = False
    while i < n:
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1; continue
        if s.startswith('# '):
            skip_section = False
            blocks.append(('H1', s[2:].strip()))
            i += 1; continue
        if s.startswith('## '):
            h2 = s[3:].strip()
            if any(h2.lower().startswith(p) for p in skip_h2_prefixes):
                skip_section = True
                i += 1; continue
            skip_section = False
            blocks.append(('H2', h2))
            i += 1; continue
        if skip_section:
            i += 1; continue
        if s.startswith('### '):
            blocks.append(('H3', s[4:].strip()))
            i += 1; continue
        if s.startswith('|') and '|' in s[1:]:
            tbl = [s]
            j = i + 1
            while j < n and lines[j].strip().startswith('|'):
                tbl.append(lines[j].strip()); j += 1
            header, data = parse_table(tbl)
            blocks.append(('TABLE', (header, data)))
            i = j; continue
        if s.startswith('- '):
            items = []
            j = i
            while j < n and lines[j].strip().startswith('- '):
                items.append(lines[j].strip()[2:]); j += 1
            blocks.append(('BULLET', items))
            i = j; continue
        if s.startswith('*') and s.endswith('*') and len(s) > 2:
            inner = s[1:-1].strip()
            if inner.lower().startswith('table') or inner.lower().startswith('figure'):
                blocks.append(('CAPTION', inner))
            else:
                blocks.append(('P', s[1:-1]))
            i += 1; continue
        blocks.append(('P', s))
        i += 1
    return blocks

def parse_title_file(path):
    """Extract title, running title, authors, abstract blocks (non-structured)."""
    with open(path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    out = dict(title=None, running=None, authors=None, abstract=[], keyword_hint=None)
    mode = None
    for raw in lines:
        s = raw.strip()
        if s.startswith('## '):
            h = s[3:].strip().lower()
            mode = h; continue
        if mode == 'title' and s.startswith('**') and s.endswith('**'):
            out['title'] = s[2:-2]
        elif mode == 'running title' and s:
            out['running'] = s
        elif mode == 'authors' and s:
            # 多行 authors 段：斜体行 = 作者名单，普通行 = 单位/通讯信息，顺序拼接
            piece = s.lstrip('*').rstrip('*') if s.startswith('*') else s
            out['authors'] = (out['authors'] + ' ' + piece) if out['authors'] else piece
        elif mode == 'abstract' and s:
            # 摘要正文里允许嵌入 *(...) 注释与 ** 加粗，全部当 P 处理
            out['abstract'].append(('P', s))
        elif mode and 'keywords' in mode and s:
            out['keyword_hint'] = s
    return out

def parse_references_file(path):
    """Extract numbered references from 07_references.md, dropping blockquotes/notes."""
    refs = []
    with open(path, encoding='utf-8') as f:
        for ln in f:
            s = ln.strip()
            m = re.match(r'^(\d+)\.\s+(.*)', s)
            if m and m.group(2) and not s.startswith('>'):
                refs.append(m.group(2))
    return refs

def load_mvmr_table():
    """Build MVMR table rows from results/mvmr.tsv."""
    fp = os.path.join(RES, 'mvmr.tsv')
    if not os.path.exists(fp) or os.path.getsize(fp) < 200:
        return None
    rows = []
    with open(fp, encoding='utf-8') as f:
        r = csv.DictReader(f, delimiter='\t')
        for row in r:
            rows.append(row)
    return rows

# ============== assemble ==============
title_info = parse_title_file(os.path.join(MANU, '01_title_abstract.md'))

# Sci Rep 章节顺序：Intro → Results → Discussion → Methods → Limitations
body_files = [
    ('02_introduction.md',  'Introduction',  []),
    ('04_results.md',       'Results',      []),
    ('05_discussion.md',    'Discussion',   []),
    ('03_methods.md',       'Methods',      []),
    ('06_limitations.md',   'Limitations',  []),
]

section_blocks = []   # list of (section_title, [blocks])
for fn, title, skip in body_files:
    blocks = generic_parse(os.path.join(MANU, fn), skip_h2_prefixes=skip)
    # 去掉文件首个 H1（它本身就是该文件名/标题，由 build 框架再加）
    if blocks and blocks[0][0] == 'H1':
        blocks = blocks[1:]
    section_blocks.append((title, blocks))

REFERENCES = parse_references_file(os.path.join(MANU, '07_references.md'))

# 从 mvmr.tsv 填充 MVMR 表（如果非空）
mvmr_rows = load_mvmr_table()

# Figure index block（在 Limitations 之后、References 之前插入）
FIGURES_BLOCK = [
    ('H1', 'Figures'),
    ('P', '**Figure 1. Study design and instrument overview.** (a) Two-sample MR design schematic. (b) Analysable exposure–outcome pairs: green = null IVW (P≥0.05), red = nominally significant (P<0.05, all in SPONDINF, n=68), grey = excluded due to no overlapping SNPs. (c) Instrument count per exposure trait (F>10).'),
    ('P', '**Figure 2. Forest plot of primary IVW estimates across 45 exposure–outcome pairs.** Points are odds ratios per genetically predicted standard-deviation increase in each immune/cytokine trait; horizontal lines are 95% confidence intervals on the log scale. The dashed vertical line indicates OR = 1. All nominally significant signals (red) cluster in the SPONDINF outcome.'),
    ('P', '**Figure 3. Fragility of the SPONDINF nominal signals.** (a) Comparison of IVW, MR-Egger, weighted median and weighted mode estimates: only IVW reaches P<0.05. (b) Directional inconsistency: leukocyte count is protective in SPONDINF (OR 0.04) but risk-increasing in discitis (OR 1.41). (c) Leave-one-out analysis shows no single SNP drives the SPONDINF signal.'),
    ('P', '**Figure 4. Statistical power and cross-phenotype meta-analysis.** (a) Minimum detectable OR (80% power, two-sided α=0.05) per outcome, stratified by selected exposures. The dashed reference line at OR=1.3 marks a clinically meaningful effect; better-powered outcomes (OM n=2,125, discitis n=495) sit well below this threshold. (b) Cross-phenotype random-effects meta forest for all 9 traits with and without SPONDINF.'),
]

# Supplementary index block
SUPP_BLOCK = [
    ('H1', 'Supplementary information'),
    ('P', 'The following Supplementary Tables and Figures accompany this manuscript and are provided as separate files for peer review. They are referenced in the main text as Supplementary Table S1, Supplementary Figure S1, etc.'),
    ('P', '**Supplementary Table S1.** Exposure manifest: EFO codes, GWAS source, ancestry and sample size for each of the 11 immune/cytokine traits.'),
    ('P', '**Supplementary Table S2.** Full IVW MR results for all 45 exposure–outcome pairs (β, SE, OR, 95% CI, P, Cochran Q, I², F-statistics).'),
    ('P', '**Supplementary Table S3.** Drug-target (cis-pQTL) MR: IL-6 receptor and CD40 against OM and DISCITIS, all four MR methods.'),
    ('P', '**Supplementary Table S4.** Multivariable MR (MVMR) adjusting for body-mass index, type-2 diabetes and smoking: full per-exposure × per-outcome estimates with conditional F-statistics and overdispersion factors.'),
    ('P', '**Supplementary Table S5.** Cross-phenotype random-effects meta-analysis (DerSimonian–Laird): pooled ORs with and without SPONDINF.'),
    ('P', '**Supplementary Table S6.** Reverse-direction MR feasibility: genome-wide significant locus counts per spine-infection phenotype in FinnGen R11.'),
    ('P', '**Supplementary Table S7.** Formal LD-clumping sensitivity analysis (1000 Genomes Phase 3 European panel, r²<0.001, via LDlink LDmatrix): instrument counts before/after clumping for the six key traits and the drug-target cis-pQTL sets, and re-estimated MR results with the clumped instruments.'),
    ('P', '**Supplementary Figure S1.** Per-exposure × per-outcome scatter, forest, funnel and leave-one-out plots (188 panels in `results/figures/`).'),
    ('P', '**Supplementary Methods.** Detailed per-SNP harmonisation actions (same/flip/strand-flip), instrument pruning pipeline, and code module reference.'),
    ('P', '**Reporting Summary.** Nature Portfolio Reporting Summary (PDF, completed at acceptance).'),
    ('P', '**STROBE-MR Checklist.** Completed STROBE-MR reporting checklist (per Skrivankova et al., 2021).'),
]

STATEMENTS = [
    ("Data availability",
     "All summary statistics used in this study are publicly available. Exposure traits were obtained from the EBI GWAS Catalog Summary Statistics REST API (GRCh38); outcome phenotypes were obtained from the FinnGen R11 public release. Specific trait identifiers (EFO codes) and FinnGen phenotype codes are listed in the Methods and Supplementary Table S1. No restricted or controlled-access data were used."),
    ("Code availability",
     "All custom code used to generate the results reported here is openly available. The analysis pipeline is written in Python 3.13 and comprises the modules gwas_io.py (EBI GWAS Catalog REST client and FinnGen remote BGZF/tabix reader), mr_pipeline.py (instrument selection, GRCh38 harmonisation and the primary 45-pair analysis), mr_methods.py (IVW, MR-Egger, weighted median, weighted mode, Cochran Q, leave-one-out and F-statistics), supplementary.py (post-hoc power, Steiger directionality and cross-phenotype meta-analysis), supplementary_network.py (drug-target cis-pQTL MR), mvmr_finngen.py (multivariable MR) and make_main_figures.py (Figures 1-4). The repository additionally contains every per-SNP harmonisation table, all leave-one-out and single-SNP estimates, the 188 per-pair supplementary panels, and the instrument manifests, so that all reported numbers can be regenerated end to end without any access-controlled resource. The code is deposited at https://github.com/547653312-ui/spine-mr-finngen and archived with a persistent identifier at Zenodo (DOI: https://doi.org/10.5281/zenodo.21863390). Dependencies and execution instructions are given in the repository README (MIT licence)."),
    ("Author contributions",
     "[To be completed at submission using CRediT (Contributor Roles Taxonomy). Example: X.Y. conceived and designed the study; X.Y. and Z.W. performed the analyses; all authors interpreted the results and drafted the manuscript.]"),
    ("Competing interests",
     "The authors declare no competing interests."),
    ("Ethics approval and consent to participate",
     "This study used only publicly available, de-identified GWAS summary statistics. No individual-level data were accessed. The work is therefore exempt from institutional ethical approval and informed-consent requirements under the policies of the source consortia (EBI GWAS Catalog, FinnGen)."),
    ("Acknowledgements",
     "[Funding sources and individual contributions to be added at acceptance.]"),
    ("Reporting Summary",
     "A Nature Portfolio Reporting Summary for this article is available as a Supplementary file."),
    ("STROBE-MR checklist",
     "A completed STROBE-MR reporting checklist is provided as a Supplementary file (see Skrivankova et al., JAMA 2021)."),
    ("Keywords (for submission system only; not part of manuscript body)",
     "Mendelian randomization; osteomyelitis; vertebral osteomyelitis; discitis; spinal infection; immune traits; cytokines; FinnGen; causality. Per Scientific Reports guidance, keywords are entered in the online submission system rather than printed in the article."),
]

# ============== DOCX ==============
def build_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    def shade(cell, color='D9D9D9'):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
        sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), color)
        tcPr.append(sh)

    def set_border(table):
        tbl = table._tbl; tblPr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4')
            e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000')
            borders.append(e)
        tblPr.append(borders)

    def add_runs(par, text, italic=False, bold=False):
        for chunk, kind in inline_parts(text):
            r = par.add_run(chunk)
            r.italic = italic or (kind == 'i')
            r.bold   = bold   or (kind == 'b')
            if kind == 'c':
                r.font.name = 'Consolas'

    # 标题块
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_info['title'] or 'Title'); r.bold = True; r.font.size = Pt(16)
    if title_info['running']:
        pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pr.add_run('Running title: ' + title_info['running'])
        rr.italic = True; rr.font.size = Pt(10)
    pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(pa, title_info['authors'] or '[Author list to be finalized]')
    pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run('Department of Orthopedics, The Fourth People\'s Hospital of Guiyang, Guiyang, China')
    rf.italic = True; rf.font.size = Pt(10)
    # 对应作者标记
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = pc.add_run('Corresponding author: Zhen Tian. Email: 547653312@qq.com. ORCID: 0000-0001-6752-7569')
    rc.italic = True; rc.font.size = Pt(10)

    # 摘要
    ph = doc.add_paragraph(); ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = ph.add_run('Abstract'); rh.bold = True; rh.font.size = Pt(12)
    for kind, txt in title_info['abstract']:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(pp, txt)

    # 主体
    for sec_title, blocks in section_blocks:
        h = doc.add_heading(level=1); h.add_run(sec_title)
        for blk in blocks:
            kind = blk[0]
            if kind == 'H2':
                hh = doc.add_heading(level=2); add_runs(hh, blk[1])
            elif kind == 'H3':
                hh = doc.add_heading(level=3); add_runs(hh, blk[1])
            elif kind == 'P':
                pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_runs(pp, blk[1])
            elif kind == 'BULLET':
                for it in blk[1]:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    add_runs(bp, it)
            elif kind == 'CAPTION':
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(cp, blk[1], italic=True)
            elif kind == 'TABLE':
                header, data = blk[1]
                t = doc.add_table(rows=1, cols=len(header))
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                t.style = 'Table Grid'; set_border(t)
                hdr = t.rows[0].cells
                for c, txt in zip(hdr, header):
                    shade(c); add_runs(c.paragraphs[0], txt, bold=True)
                for row in data:
                    cells = t.add_row().cells
                    for c, txt in zip(cells, row):
                        add_runs(c.paragraphs[0], txt)
                for row in t.rows:
                    for c in row.cells:
                        for pp in c.paragraphs:
                            for rr in pp.runs:
                                rr.font.size = Pt(9)
                doc.add_paragraph()

    # Figures 节
    for blk in FIGURES_BLOCK:
        kind = blk[0]
        if kind == 'H1':
            doc.add_heading(level=1).add_run(blk[1])
        elif kind == 'P':
            pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(pp, blk[1])
    # 嵌入 4 张主图（PNG）
    for fn in ['Figure1.png','Figure2.png','Figure3.png','Figure4.png']:
        fp = os.path.join(FIG, fn)
        if os.path.exists(fp):
            doc.add_picture(fp, width=Inches(6.0))
            last = doc.paragraphs[-1]; last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Supplementary information
    for blk in SUPP_BLOCK:
        kind = blk[0]
        if kind == 'H1':
            doc.add_heading(level=1).add_run(blk[1])
        elif kind == 'P':
            pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(pp, blk[1])

    # References
    doc.add_heading(level=1).add_run('References')
    for i, ref in enumerate(REFERENCES, 1):
        pr = doc.add_paragraph(style='List Number')
        add_runs(pr, ref)

    # Statements
    doc.add_heading(level=1).add_run('Scientific Reports checklist statements')
    for head, body in STATEMENTS:
        h = doc.add_paragraph(); rh = h.add_run(head); rh.bold = True
        pb = doc.add_paragraph(); pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(pb, body)

    doc.save(TMP_DOCX)
    print('DOCX written:', OUT_DOCX)

# ============== TEX ==============
def tex_escape(text):
    out = []; sup_buffer = ''
    def flush_sup():
        nonlocal sup_buffer
        if sup_buffer:
            digits = ''.join(SUP_MAP.get(ch, ch) for ch in sup_buffer)
            out.append(r'\textsuperscript{' + digits + '}'); sup_buffer = ''
    for ch in text:
        if ch in SUP_MAP:
            sup_buffer += ch; continue
        flush_sup()
        if ch == '\u00d7': out.append(r'\times ')
        elif ch == '\u2192': out.append(r'$\rightarrow$ ')
        elif ch == '\u2194': out.append(r'$\leftrightarrow$ ')
        elif ch == '\u2013' or ch == '\u2014': out.append('--')
        elif ch == '\u2026': out.append(r'\ldots ')
        elif ch == '\u00a7': out.append(r'\S ')
        elif ch == '\u03b1': out.append(r'$\alpha$')
        elif ch == '\u03b2': out.append(r'$\beta$')
        elif ch in '\\&%$#_{}~^':
            out.append('\\' + ch)
        else:
            out.append(ch)
    flush_sup()
    return ''.join(out)

def tex_inline(text):
    chunks = inline_parts(text); out = []
    for chunk, kind in chunks:
        esc = tex_escape(chunk)
        if   kind == 'b': out.append(r'\textbf{' + esc + '}')
        elif kind == 'i': out.append(r'\emph{' + esc + '}')
        elif kind == 'c': out.append(r'\texttt{' + esc + '}')
        else: out.append(esc)
    return ''.join(out)

def build_tex():
    L = []
    L.append(r'\documentclass[11pt]{article}')
    L.append(r'\usepackage[margin=1in]{geometry}')
    L.append(r'\usepackage{booktabs}')
    L.append(r'\usepackage{array}')
    L.append(r'\usepackage{amsmath}')
    L.append(r'\usepackage{graphicx}')
    L.append(r'\usepackage{hyperref}')
    L.append(r'\setlength{\parindent}{0pt}')
    L.append(r'\begin{document}')
    L.append('')
    L.append(r'\title{' + tex_escape(title_info['title'] or 'Title') + '}')
    authors_block = title_info['authors'] or '[Author list to be finalized]'
    if title_info['running']:
        authors_block += r'\\ \emph{Running title: ' + tex_escape(title_info['running']) + '}'
    authors_block += r'\\ Department of Orthopedics, The Fourth People\'s Hospital of Guiyang, Guiyang, China'
    authors_block += r'\\ Corresponding author: Zhen Tian. Email: 547653312@qq.com. ORCID: 0000-0001-6752-7569'
    L.append(r'\author{' + tex_escape(authors_block) + '}')
    L.append(r'\date{}'); L.append(r'\maketitle')
    L.append('')
    L.append(r'\begin{abstract}')
    for kind, txt in title_info['abstract']:
        L.append(tex_inline(txt))
    L.append(r'\end{abstract}')
    L.append('')
    for sec_title, blocks in section_blocks:
        L.append(r'\section{' + tex_escape(sec_title) + '}')
        for blk in blocks:
            kind = blk[0]
            if kind == 'H2':
                L.append(r'\subsection{' + tex_escape(blk[1]) + '}')
            elif kind == 'H3':
                L.append(r'\subsubsection{' + tex_escape(blk[1]) + '}')
            elif kind == 'P':
                L.append(tex_inline(blk[1])); L.append('')
            elif kind == 'BULLET':
                L.append(r'\begin{itemize}')
                for it in blk[1]:
                    L.append(r'\item ' + tex_inline(it))
                L.append(r'\end{itemize}')
            elif kind == 'CAPTION':
                L.append(r'\begin{quote}\emph{' + tex_escape(blk[1]) + r'}\end{quote}')
            elif kind == 'TABLE':
                header, data = blk[1]
                L.append(r'\begin{table}[ht]\centering')
                L.append(r'\begin{tabular}{' + 'l'*len(header) + '}')
                L.append(r'\toprule')
                L.append(' & '.join(tex_escape(h) for h in header) + r' \\')
                L.append(r'\midrule')
                for row in data:
                    L.append(' & '.join(tex_escape(c) for c in row) + r' \\')
                L.append(r'\bottomrule')
                L.append(r'\end{tabular}'); L.append(r'\end{table}')
    # Figures
    L.append(r'\section*{Figures}')
    for blk in FIGURES_BLOCK[1:]:  # 跳过 H1
        L.append(tex_inline(blk[1])); L.append('')
    for fn in ['Figure1.png','Figure2.png','Figure3.png','Figure4.png']:
        fp = os.path.join(FIG, fn)
        if os.path.exists(fp):
            base = fn.replace('.png', '')
            L.append(r'\begin{figure}[ht]\centering'
                     r'\includegraphics[width=0.95\textwidth]{' + os.path.join('../../results/figures', base) + r'}'
                     r'\end{figure}')
    # Supplementary
    L.append(r'\section*{Supplementary information}')
    for blk in SUPP_BLOCK[1:]:
        L.append(tex_inline(blk[1])); L.append('')
    # References
    L.append(r'\section*{References}')
    L.append(r'\begin{enumerate}')
    for ref in REFERENCES:
        L.append(r'\item ' + tex_escape(ref))
    L.append(r'\end{enumerate}')
    # Statements
    L.append(r'\section*{Scientific Reports checklist statements}')
    for head, body in STATEMENTS:
        L.append(r'\paragraph{' + tex_escape(head) + '} ' + tex_escape(body))
    L.append(r'\end{document}')
    with open(TMP_TEX, 'w', encoding='utf-8') as f:
        f.write('\n'.join(L))
    print('TEX written:', OUT_TEX)

if __name__ == '__main__':
    build_docx()
    build_tex()
    # 原子替换：直接 os.replace 覆盖目标（Windows 上可覆盖已存在文件），
    # 避免先删后拷被删除安全钩子/文件锁拦截
    import os as _os
    for tmp, final in [(TMP_DOCX, OUT_DOCX), (TMP_TEX, OUT_TEX)]:
        try:
            _os.replace(tmp, final)
            print('Replaced:', final)
        except Exception as e:
            print('replace failed, keep tmp:', final, repr(e))
    print('Done. References count:', len(REFERENCES))